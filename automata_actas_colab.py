# -*- coding: utf-8 -*-
"""
=====================================================================
DATAENG - AUTOMATIZADOR DE ACTAS INSTITUCIONALES (VERSIÓN GOOGLE COLAB)
=====================================================================

Este script permite generar actas de reuniones de forma automatizada
utilizando Inteligencia Artificial (Gemini 2.5 Flash) a partir de una
transcripción y una plantilla de Word (.docx) con marcadores.

---------------------------------------------------------------------
GUÍA DE USO PARA EL USUARIO (SECRETARÍA / ASISTENTES)
---------------------------------------------------------------------

1. CÓMO OBTENER LA CLAVE DE API DE GEMINI (GRATUITA)
   La clave de API (API Key) es necesaria para que el programa pueda
   "hablar" con la Inteligencia Artificial de Google.
   a) Ingresa a Google AI Studio: https://aistudio.google.com/
   b) Inicia sesión con cualquier cuenta de Gmail (@gmail.com).
   c) Haz clic en el botón azul "Get API Key" (Obtener clave de API).
   d) Luego haz clic en "Create API Key" (Crear clave de API).
   e) Selecciona un proyecto por defecto y haz clic en "Create API Key in existing project".
   f) Copia la clave generada (es una serie larga de letras y números)
      y guárdala en un lugar seguro. ¡No la compartas públicamente!

2. CÓMO PREPARAR LA PLANTILLA DE WORD (.docx)
   El programa busca palabras especiales llamadas "marcadores" dentro de
   tu plantilla de Word y las reemplaza con la información de la reunión.
   Los marcadores deben estar escritos entre signos menores/mayores.
   Puedes usar el formato simple <Marcador> o el formato doble <<Marcador>>.
   
   Ejemplo de marcadores recomendados en la plantilla:
   - <<FechaActual>> o <FechaActual>: Para la fecha de la reunión.
   - <<TemaReunion>> o <TemaReunion>: El título o asunto principal.
   - <<Agenda1>>, <<Agenda2>>, <<Agenda3>>...: Para los puntos del orden del día.
   - <<Comentario1>>, <<Comentario2>>, <<Comentario3>>...: Para el desarrollo de cada tema.
   - <<Responsable1>>, <<Fecha1>>, <<Accion1>>, <<Tipo1>>...: Para los compromisos asignados.
   - <<Tipo1>>, <<Tipo2>>...: Clasificación automática (Informativo, Decisorio o Estratégico).
   - <<ResponsablesVarios>>, <<AccionVarios>>, <<TipoVarios>>, <<FechaVarios>>: Para compromisos varios.

3. FORMATOS DE TRANSCRIPCIÓN Y RECOMENDACIÓN
   El archivo de transcripción contiene todo lo que se habló en la reunión.
   - FORMATO RECOMENDADO: Texto plano (.txt). Es el más liviano, rápido de
     procesar y evita errores de lectura.
   - Otros formatos soportados: Documento de Word (.docx) y subtítulos WebVTT (.vtt).
   - ¿Cómo obtenerlo?: Puedes descargar la transcripción de tus reuniones en
     Microsoft Teams, Zoom o Google Meet y guardarla en tu computador en
     formato de texto (.txt).

4. QUÉ HACER CUANDO TE COMPARTAN ESTE CUADERNO DE GOOGLE COLAB
   a) Abre el enlace de Google Colab que te compartieron.
   b) Recomendado: Haz una copia en tu propio Google Drive haciendo clic en
      "Archivo" (File) -> "Guardar una copia en Drive" (Save a copy in Drive).
   c) En el panel derecho de la celda de configuración, verás dos campos:
      - "API_KEY": Pega allí la clave que obtuviste en el paso 1.
      - "NIVEL_DETALLE": Selecciona qué tan detallado deseas el resumen
        de los comentarios ("Resumido", "Medio" o "Detallado").
   d) Haz clic en el botón de reproducción (el círculo con un triángulo "Play")
      situado a la izquierda de la celda de código.
   e) El programa instalará automáticamente todo lo que necesita.
   f) Te aparecerá un botón que dice "Elegir archivos" (Choose files). Haz clic
      allí y selecciona tu PLANTILLA de Word (.docx) desde tu computador.
   g) A continuación, aparecerá otro botón para que selecciones el archivo
      con la TRANSCRIPCIÓN (.txt, .docx o .vtt).
   h) Espera unos segundos mientras la Inteligencia Artificial procesa la reunión.
   i) ¡Listo! El acta final se descargará automáticamente en tu carpeta de
      "Descargas" (Downloads) con el nombre "Acta_Generada_...docx".
"""

import os
import re
import sys
import json
import time
import subprocess

# =====================================================================
# 0. AUTOINSTALACIÓN DE DEPENDENCIAS (Solo si no están presentes)
# =====================================================================
def instalar_dependencias():
    librerias_faltantes = []
    
    try:
        import docx
    except ImportError:
        librerias_faltantes.append("python-docx")
        
    try:
        from google import genai
    except ImportError:
        librerias_faltantes.append("google-genai")
        
    if librerias_faltantes:
        print(f"Instalando librerías necesarias: {', '.join(librerias_faltantes)}...")
        try:
            # Ejecutar instalación silenciosa
            subprocess.check_call([sys.executable, "-m", "pip", "install", *librerias_faltantes], 
                                  stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            print("¡Instalación de dependencias completada con éxito!\n")
        except Exception as e:
            print(f"Error al instalar dependencias de forma automática: {e}")
            print("Por favor, ejecuta manualmente: !pip install google-genai python-docx")

# Ejecutar la autoinstalación al importar o iniciar el script
instalar_dependencias()

# Importaciones requeridas después de asegurar su instalación
from docx import Document
from google import genai
from google.genai import types

# =====================================================================
# 1. DETECCIÓN DEL ENTORNO DE EJECUCIÓN (Google Colab vs. Local)
# =====================================================================
def es_entorno_colab():
    try:
        import google.colab
        return True
    except ImportError:
        return False

# =====================================================================
# 2. INTELIGENCIA CONTEXTUAL Y LLAMADO A GEMINI
# =====================================================================
def generar_mapeo_con_ia(transcripcion, lista_marcadores, api_key, nivel_detalle):
    if not api_key:
        raise ValueError("La clave de API de Gemini está vacía. Por favor, ingrésala en la sección de configuración.")
        
    client = genai.Client(api_key=api_key)
    
    # Crear un mapeo para limpiar tildes y caracteres especiales en las llaves del JSON
    def simplificar_llave(texto):
        remplazos = {"á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u", "ñ": "n", 
                     "Á": "A", "É": "E", "Í": "I", "Ó": "O", "Ú": "U", "Ñ": "N"}
        for k, v in remplazos.items():
            texto = texto.replace(k, v)
        return re.sub(r'[^a-zA-Z0-9_]', '', texto)

    # Evitar colisiones de llaves al simplificarlas
    mapeo_llaves = {}
    llaves_usadas = set()
    for m in lista_marcadores:
        llave_base = simplificar_llave(m)
        llave_limpia = llave_base
        contador = 1
        while llave_limpia in llaves_usadas:
            llave_limpia = f"{llave_base}_{contador}"
            contador += 1
        llaves_usadas.add(llave_limpia)
        mapeo_llaves[llave_limpia] = m
        
    llaves_limpias = list(mapeo_llaves.keys())
    
    # Construcción del esquema estructurado
    propiedades_dinamicas = {
        llave: types.Schema(type="STRING", description=f"Datos para el campo {mapeo_llaves[llave]}") 
        for llave in llaves_limpias
    }
    
    esquema_dinamico = types.Schema(
        type="OBJECT",
        properties=propiedades_dinamicas,
        required=llaves_limpias
    )

    if nivel_detalle == "Resumido":
        instrucciones_longitud = "- SÉ CONCISO Y DIRECTO (máximo 1 o 2 oraciones breves para bloques de desarrollo)."
        temp = 0.15
    elif nivel_detalle == "Medio":
        instrucciones_longitud = "- Redacta párrafos profesionales balanceados (de 2 a 3 líneas de extensión para bloques de desarrollo)."
        temp = 0.25
    else: # Detallado
        instrucciones_longitud = """
        - Desarrolla una prosa institucional formal, muy amplia, exhaustiva y altamente descriptiva para bloques de desarrollo.
        - Cada sección analizada de este tipo debe contar con un párrafo extenso (mínimo de 4 a 6 líneas de extensión).
        """
        temp = 0.35

    prompt = f"""
    Actúas como un Ingeniero de Datos experto en procesamiento de lenguaje natural y redacción corporativa institucional. 
    Analiza la transcripción de la reunión adjunta y extrae la información requerida para los campos de la plantilla.
    
    Reglas críticas de formato y contenido:
    1. CONTENIDO LIMPIO (SIN MARCADORES): Ningún valor (value) debe contener los caracteres '<<' o '>>'.
    2. PRESERVACIÓN DE DATOS PUROS: No incluyas prefijos de numeración (como '1. ', '2. ') en las Agendas o Comentarios. El JSON guarda solo texto puro.
    3. SÍNTESIS NARRATIVA FORMAL: Los campos referentes a debates, resúmenes o comentarios de la agenda (ej: Comentario1 a Comentario5, ComentarioVarios) deben ser descriptivos, amplios y redactados en prosa institucional.
    4. ACCIONES CON VERBO EN INFINITIVO: Cada campo que defina una acción o tarea pendiente (ej: Accion1 a Accion5) debe iniciar obligatoriamente con un verbo en infinitivo (ej: Gestionar, Implementar, Diseñar, Validar).
    5. FIDELIDAD TOTAL: Datos exactos, cifras, nombres y fechas límites deben ser extraídos fielmente sin inventar información. Si un dato no se menciona en absoluto en la transcripción, pon un guión (-) o déjalo vacío.
    6. REGLAS DE EXTENSIÓN PARA CAMPOS DE DESARROLLO (TIPO C) - Nivel actual: {nivel_detalle}:
       {instrucciones_longitud}
    7. REGLA DE DISTRIBUCIÓN Y SÍNTESIS DE TEMAS (Especial para campos numerados de 1 a 5):
       - Identifica los temas o puntos clave discutidos en la reunión (hasta 5 temas principales).
       - Distribuye estos temas de forma secuencial en los campos numerados (Agenda1 a Agenda5, Comentario1 a Comentario5, Accion1 a Accion5).
       - Si la reunión tuvo menos de 5 temas diferenciados, agrupa los temas identificados empezando desde el 1 y deja los restantes (por ejemplo, Agenda4, Agenda5 si solo hubo 3 temas) en blanco o con un guion (-), junto con sus respectivos Comentarios, Acciones, Responsables, Tipos y Fechas.
       - Cada punto de la Agenda numerado (ej: Agenda1) debe corresponderse temáticamente con su Comentario (ej: Comentario1) y con su Acción de seguimiento (ej: Accion1, Tipo1, Responsable1, Fecha1) si se acordó algún compromiso para ese tema.
    
    Transcripción de la reunión:
    \"\"\"{transcripcion}\"\"\"
    """
    
    max_intentos = 5
    espera = 3
    for intento in range(max_intentos):
        try:
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=esquema_dinamico,
                    temperature=temp,
                    max_output_tokens=8000
                ),
            )
            break
        except Exception as e:
            if intento == max_intentos - 1:
                raise e
            print(f"Intento {intento+1} de Gemini falló. Reintentando en {espera} segundos...")
            time.sleep(espera)
            espera *= 2
            
    datos_ia = json.loads(response.text)
    datos_restaurados = {mapeo_llaves[llave_limpia]: valor for llave_limpia, valor in datos_ia.items()}
    return datos_restaurados

# =====================================================================
# 3. EXTRACCIÓN Y REEMPLAZO DIRECTO EN EL DOCUMENTO WORD
# =====================================================================
def extraer_marcadores_word(ruta_plantilla):
    doc = Document(ruta_plantilla)
    marcadores = set()
    patron = r'(?:<<|<)([a-zA-Z0-9_\-\s\?¿\(\)\.]+)(?:>>|>)'
    
    # Buscar en párrafos comunes
    for p in doc.paragraphs:
        encontrados = re.findall(patron, p.text)
        for e in encontrados:
            marcadores.add(e.strip())
            
    # Buscar dentro de tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    encontrados = re.findall(patron, p.text)
                    for e in encontrados:
                        marcadores.add(e.strip())
                        
    return list(marcadores)

def generar_nuevo_documento(ruta_plantilla, ruta_salida, datos_mapeados):
    doc = Document(ruta_plantilla)
    
    def reemplazar_marcando(texto_original, datos):
        texto_modificado = texto_original
        for k, v in datos.items():
            marcador_doble = f"<<{k}>>"
            marcador_simple = f"<{k}>"
            
            if marcador_doble in texto_modificado:
                texto_modificado = texto_modificado.replace(marcador_doble, str(v))
            if marcador_simple in texto_modificado:
                texto_modificado = texto_modificado.replace(marcador_simple, str(v))
        return texto_modificado

    # 1. Párrafos fuera de tablas
    for p in doc.paragraphs:
        if "<<" in p.text or "<" in p.text:
            p.text = reemplazar_marcando(p.text, datos_mapeados)
            
    # 2. Celdas dentro de tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    if "<<" in p.text or "<" in p.text:
                        p.text = reemplazar_marcando(p.text, datos_mapeados)
                    
    doc.save(ruta_salida)

# Helper para leer archivos de transcripción de forma robusta
def leer_archivo_transcripcion(ruta):
    if ruta.endswith('.docx'):
        doc = Document(ruta)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        with open(ruta, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

# =====================================================================
# 4. FUNCIÓN PRINCIPAL DE EJECUCIÓN (PIPELINE)
# =====================================================================
def ejecutar_automatizacion(api_key, nivel_detalle, ruta_plantilla=None, ruta_transcripcion=None):
    """
    Ejecuta el pipeline completo de automatización de actas.
    Si las rutas son None, se manejará la subida interactiva según el entorno.
    """
    colab = es_entorno_colab()
    
    if not api_key or api_key.strip() == "":
        print("❌ ERROR: La clave de API de Gemini está vacía.")
        print("Obtén una clave gratuita en: https://aistudio.google.com/ e introdúcela en el campo 'API_KEY'.")
        return
        
    print("--- AUTOMATIZADOR DE ACTAS DATAENG ---")
    print(f"Modo detectado: {'Google Colab' if colab else 'Ejecución Local'}")
    print(f"Nivel de detalle solicitado: {nivel_detalle}\n")
    
    # 1. Cargar archivos
    if colab:
        from google.colab import files
        
        # Subida de la plantilla
        if not ruta_plantilla:
            print("1. SUBIDA DE LA PLANTILLA WORD:")
            print("Por favor, selecciona tu plantilla institucional (.docx)...")
            subidos = files.upload()
            if not subidos:
                print("❌ ERROR: No se subió ninguna plantilla.")
                return
            ruta_plantilla = list(subidos.keys())[0]
            print(f"Plantilla cargada: {ruta_plantilla}\n")
            
        # Subida de la transcripción
        if not ruta_transcripcion:
            print("2. SUBIDA DE LA TRANSCRIPCIÓN DE LA REUNIÓN:")
            print("Selecciona tu archivo de transcripción (.txt, .docx o .vtt)...")
            subidos = files.upload()
            if not subidos:
                print("❌ ERROR: No se subió ninguna transcripción.")
                return
            ruta_transcripcion = list(subidos.keys())[0]
            print(f"Transcripción cargada: {ruta_transcripcion}\n")
    else:
        # Modo Consola Local interactiva si las rutas no se proveen
        if not ruta_plantilla:
            ruta_plantilla = input("Ingresa la ruta de la Plantilla Word (.docx) [Ej: Formato_F-2-2-16_Place.docx]: ").strip()
        if not ruta_transcripcion:
            ruta_transcripcion = input("Ingresa la ruta de la Transcripción (.txt, .docx, .vtt) [Ej: Transcripcion_Chipaque.txt]: ").strip()
            
        if not os.path.exists(ruta_plantilla):
            print(f"❌ ERROR: La ruta de la plantilla no existe: {ruta_plantilla}")
            return
        if not os.path.exists(ruta_transcripcion):
            print(f"❌ ERROR: La ruta de la transcripción no existe: {ruta_transcripcion}")
            return

    # 2. Análisis del documento
    try:
        print("🔍 Analizando marcadores en la plantilla...")
        marcadores = extraer_marcadores_word(ruta_plantilla)
        if not marcadores:
            print("❌ ERROR: No se encontraron marcadores válidos (ej: <<Marcador>> o <Marcador>) en la plantilla.")
            return
        print(f"Marcadores identificados ({len(marcadores)}): {', '.join(marcadores)}")
        
        # 3. Lectura de la transcripción
        print("\n📖 Leyendo transcripción...")
        texto_transcripcion = leer_archivo_transcripcion(ruta_transcripcion)
        if not texto_transcripcion.strip():
            print("❌ ERROR: El archivo de transcripción está vacío.")
            return
            
        # 4. Procesamiento IA con Gemini
        print("\n🤖 Procesando transcripción con Gemini 2.5 Flash...")
        print("Esto puede demorar unos segundos debido al análisis del contenido...")
        datos_mapeados = generar_mapeo_con_ia(texto_transcripcion, marcadores, api_key, nivel_detalle)
        
        # 5. Generación del acta final
        print("\n✍️ Insertando información y generando el documento definitivo...")
        nombre_salida = f"Acta_Generada_{nivel_detalle}.docx"
        
        if colab:
            ruta_salida = nombre_salida
        else:
            # Localmente se guarda en la carpeta de la plantilla
            directorio = os.path.dirname(os.path.abspath(ruta_plantilla))
            ruta_salida = os.path.join(directorio, nombre_salida)
            
        generar_nuevo_documento(ruta_plantilla, ruta_salida, datos_mapeados)
        print(f"✅ ¡Acta generada con éxito! Archivo creado: {ruta_salida}")
        
        # 6. Descarga del archivo si está en Colab
        if colab:
            print("\n📥 Iniciando descarga automática del acta generada en tu computador...")
            files.download(ruta_salida)
            print("¡Descarga completada! Revisa tu carpeta de descargas.")
            
    except Exception as e:
        print(f"\n❌ Ocurrió un error crítico durante el procesamiento: {str(e)}")


# =====================================================================
# BLOQUE DE EJECUCIÓN DIRECTA PARA GOOGLE COLAB / JUPYTER FORMS
# =====================================================================
# Las variables de abajo con #@param permiten crear formularios interactivos en Colab.

if __name__ == "__main__" and es_entorno_colab():
    #@title Configuración del Automatizador de Actas { display-mode: "form" }
    API_KEY = "" #@param {type:"string"}
    NIVEL_DETALLE = "Medio" #@param ["Resumido", "Medio", "Detallado"]
    
    if API_KEY.strip() == "":
        print("👉 Por favor, ingresa tu API KEY de Gemini en el formulario de la derecha y ejecuta la celda.")
    else:
        ejecutar_automatizacion(api_key=API_KEY, nivel_detalle=NIVEL_DETALLE)

elif __name__ == "__main__":
    # Modo de consola normal para desarrollo o prueba local
    print("--- Modo Consola Local ---")
    api_key_env = ""
    # Intentar cargar config.json si existe para comodidad en desarrollo local
    if os.path.exists("config.json"):
        try:
            with open("config.json", "r", encoding="utf-8") as f:
                api_key_env = json.load(f).get("api_key", "")
        except:
            pass
            
    api_key = input(f"Ingresa tu Gemini API Key [{api_key_env[:6]}...]: ").strip() if not api_key_env else input("Ingresa tu Gemini API Key: ").strip()
    if not api_key and api_key_env:
        api_key = api_key_env
        
    detalle = input("Nivel de detalle (Resumido/Medio/Detallado) [Medio]: ").strip()
    if detalle not in ["Resumido", "Medio", "Detallado"]:
        detalle = "Medio"
        
    ejecutar_automatizacion(api_key=api_key, nivel_detalle=detalle)
